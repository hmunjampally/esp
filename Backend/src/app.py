from flask import Flask, request, jsonify, make_response
from flask_cors import CORS
from docx import Document
from pathlib import Path
import logging
import os
import json

def create_app():
    app = Flask(__name__)
    CORS(app)

    uploads_dir = os.path.join(app.instance_path, 'uploads')
    os.makedirs(uploads_dir, exist_ok=True)

    @app.route('/upload', methods=['POST'])
    def file_upload():
        if 'file' not in request.files:
            logging.error('No file part in request')
            return {'message': 'No file part'}, 400
        
        file = request.files['file']
        
        if file.filename == '':
            logging.error('No selected file')
            return {'message': 'No selected file'}, 400
        
        if file and file.mimetype == 'text/plain':
            filename = os.path.join(uploads_dir, file.filename)
            file.save(filename)
            return {'message': 'Success', 'filename':filename}, 200  
        else:
            logging.error('Invalid file type')
            return {'message': 'Invalid file type'}, 400

    @app.route('/api/data', methods=['GET'])
    def get_data():
        data = {
            'options': [
                {'label': 'Option 1', 'value': 1},
                {'label': 'Option 2', 'value': 2},
                {'label': 'Option 3', 'value': 3}
            ]
        }
        return jsonify(data), 200

    @app.route('/find-sections', methods=['POST'])
    def find_sections():
        # Check for missing request data
        if 'file_path' not in request.json:
            return 'Missing file_path field', 400
        if 'search_terms' not in request.json:
            return 'Missing search_terms field', 400
        if 'sections' not in request.json:
            return 'Missing sections field', 400
        if 'specifyLines' not in request.json:
            return 'Missing specifyLines field', 400
        if 'use_total_lines' not in request.json:
            return 'Missing use_total_lines field', 400
        if 'total_lines' not in request.json:
            return 'Missing total_lines field', 400

        # Get the data from the request
        file_path = request.json['file_path']
        search_terms = request.json['search_terms']
        sections_string = request.json['sections']
        sections = json.loads(sections_string)
        tempSpecifyLines = request.json['specifyLines']
        use_total_lines = bool(request.json['use_total_lines'])
        total_lines = int(request.json['total_lines'])

        specifyLines = []
        num = 0
        while num <= len(sections):
            specifyLines.append(tempSpecifyLines)
            num += num + 1

        # Read the file
        with open(file_path, 'r') as f:
            Lines = f.readlines()

        # Create a new document
        document = Document()

        # Iterate over the search terms and find the corresponding sections
        for term in search_terms:
            lineNo = 0
            termLineNo = []
            termsNum = 0
            for line in Lines:
                if term in line:
                    termLineNo.append(lineNo)
                    termsNum += 1
                lineNo += 1

        # Add the sections to the document
        for i in sections:
            section_lines = specifyLines[i-1].split()
            start_line = termLineNo[i-1]
            line_empty = 0

            if section_lines[0].upper() == 'WHOLE' and not use_total_lines:
                while line_empty == 0:
                    if Lines[start_line] != "\n":
                        section = document.add_paragraph(Lines[start_line])
                        start_line += 1
                    else:
                        line_empty = 1

            if section_lines[0].upper() == 'WHOLE' and use_total_lines:
                for _ in range(total_lines - start_line + termLineNo[i-1]):
                    section = document.add_paragraph(Lines[start_line])
                    start_line += 1
                    line_empty = 1
                else:
                    start_line += 1
                    line_empty = 1

            elif section_lines[0].upper() == 'FIRST':
                line_count = -1
                while line_count < int(section_lines[1]):
                    section = document.add_paragraph(Lines[start_line])
                    start_line += 1
                    line_count += 1

            elif section_lines[0].upper() == 'LAST':
                line_count = -1
                document.add_paragraph(Lines[start_line])
                document.add_paragraph(Lines[start_line + 1])
                while line_count < int(section_lines[1]):
                    section = document.add_paragraph(Lines[start_line+10])
                    start_line += 1
                    line_count += 1

            elif section_lines[0].upper() == 'SPECIFIC':
                specific_lines = [int(l) for l in section_lines[1].split(",")]
                document.add_paragraph(Lines[start_line])
                for l in specific_lines:
                    section = document.add_paragraph(Lines[start_line + l + 1])

        # Convert the document to bytes and return it as a response
        try:
            docx_bytes = save_document_to_bytes(document)
            response = make_response(docx_bytes)
            response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response.headers.set('Content-Disposition', 'attachment', filename='output.docx')
            return response
        except Exception as e:
            return f'Error saving document: {e}', 500

    def save_document_to_bytes(document):
        """Save the Word document to a byte string."""
        from io import BytesIO
        file_stream = BytesIO()
        document.save(file_stream)
        return file_stream.getvalue()

    return app

if __name__ == "__main__":
    app = create_app()
    app.run(debug=True)
